import os
import re
import json
import requests
from flask import Flask, request, jsonify, send_file, render_template
from flask_cors import CORS
import pypdf
from bs4 import BeautifulSoup
from google import genai
from google.genai import types
from io import BytesIO
from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import ParagraphStyle
from reportlab.graphics.shapes import Drawing, Line
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

app = Flask(__name__)
CORS(app)

# ── Frontend ────────────────────────────────────────────────────────────────────
@app.route('/')
def index():
    return render_template('index.html')


# ── Utility helpers ─────────────────────────────────────────────────────────────

def scrape_url(url):
    try:
        headers = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64)'}
        resp = requests.get(url, headers=headers, timeout=10)
        resp.raise_for_status()
        soup = BeautifulSoup(resp.text, 'html.parser')
        for el in soup(["script", "style", "nav", "footer", "header", "aside"]):
            el.decompose()
        raw = soup.get_text(separator=' ')
        lines = (l.strip() for l in raw.splitlines())
        chunks = (p.strip() for l in lines for p in l.split("  "))
        return '\n'.join(c for c in chunks if c)
    except Exception as e:
        raise Exception(f"Failed to scrape URL: {e}")


def extract_pdf_text(file_stream):
    try:
        reader = pypdf.PdfReader(file_stream)
        text = ""
        for page in reader.pages:
            t = page.extract_text()
            if t:
                text += t + "\n"
        return text.strip()
    except Exception as e:
        raise Exception(f"Failed to read PDF: {e}")


def strip_markdown(text):
    """Remove **bold**, *italic*, __bold__, _italic_ markers."""
    if not text:
        return text
    text = re.sub(r'\*{1,2}([^*]+)\*{1,2}', r'\1', text)
    text = re.sub(r'_{1,2}([^_]+)_{1,2}', r'\1', text)
    return text.strip()


# ── Analyse endpoint ─────────────────────────────────────────────────────────────
@app.route('/api/analyze', methods=['POST'])
def analyze():
    try:
        if 'cv' not in request.files:
            return jsonify({"error": "Missing CV file (PDF)"}), 400
        cv_file = request.files['cv']
        if cv_file.filename == '':
            return jsonify({"error": "No selected file"}), 400

        jd_text = request.form.get('job_description_text', '').strip()
        jd_url  = request.form.get('job_description_url', '').strip()
        if not jd_text and not jd_url:
            return jsonify({"error": "Provide a job description text or URL"}), 400

        if jd_url:
            try:
                scraped = scrape_url(jd_url)
                jd_text = f"{jd_text}\n\n[Scraped]:\n{scraped}".strip()
            except Exception as e:
                return jsonify({"error": str(e)}), 500

        try:
            cv_text = extract_pdf_text(cv_file)
        except Exception as e:
            return jsonify({"error": str(e)}), 400

        if not cv_text:
            return jsonify({"error": "PDF contains no readable text."}), 400

        api_key = request.headers.get("X-Gemini-Key") or os.environ.get("GEMINI_API_KEY")
        if not api_key:
            return jsonify({"error": "Missing Gemini API Key."}), 400

        try:
            client = genai.Client(api_key=api_key)
        except Exception as e:
            return jsonify({"error": f"Gemini client error: {e}"}), 500

        prompt = f"""
You are an expert resume writer and ATS optimization specialist.
IMPORTANT: All text values in your JSON output must be plain text only.
Do NOT use any markdown formatting (no **bold**, no *italic*, no - bullets, no # headers).
Write natural sentences exactly as they would appear in a professional CV.
When incorporating keywords into sentences, use natural sentence-case capitalisation. Do NOT capitalise keywords mid-sentence unless they are genuine proper nouns (e.g. Python, Agile, Amazon) or acronyms (e.g. KPI, OKR, SQL). Common phrases like "data-driven decision making" or "stakeholder management" should remain in lower case when used mid-sentence.

Your tasks:
1. Parse the CV into structured JSON.
2. Identify 10-15 keywords or skills from the Job Description and SCORE each one using the 4 criteria below. The final priority rank (1 = most important) is determined by total score descending.
3. Propose natural sentence rewrites in the CV Profile or Experience bullets to weave in high-priority keywords. Never invent achievements — only rephrase what is already there.
4. Scan the entire CV text for redundancy: repeated words, phrases or concepts that appear more than once across sections. For each issue, identify the location and suggest a cleaner, more human-sounding alternative. Focus on varied language and avoiding repetition.

KEYWORD SCORING CRITERIA (assign each keyword a score out of 9):

A. JD Prominence (0-3):
   +3 if the keyword appears in the job title, opening paragraph, or under an explicit "Required" / "Must have" section
   +2 if it appears in the main responsibilities section
   +1 if it appears only in "Nice to have" or closing sections
   +0 if it is only mentioned once in passing

B. JD Frequency (0-2):
   +2 if the keyword (or close synonyms) appears 3 or more times in the JD
   +1 if it appears twice
   +0 if it appears only once

C. Specificity (0-1):
   +1 if it is a concrete, specific skill (tool, methodology, certification, domain)
   +0 if it is a generic soft skill (e.g. "communication", "teamwork")

D. Fit with candidate CV background (0-3):
   +3 if it directly relates to the candidate's existing experience or domain
   +2 if it is adjacent — the candidate could credibly claim this skill based on their background
   +1 if it is loosely related but requires some stretch
   +0 if it is unrelated to the candidate's background and would be implausible to include

Only recommend keywords where D score is 1 or higher. Do NOT surface keywords with D=0.
Set matching_status to "missing" if the keyword is absent from the CV, or "under-represented" if present but not prominent.

CV Text:
{cv_text}

Job Description:
{jd_text}

Return ONE JSON object matching this schema exactly:
{{
  "parsed_cv": {{
    "name": "Full Name",
    "title": "Professional Title",
    "contact": "Email | Phone | Location",
    "profile": "Profile summary paragraph",
    "experience": [
      {{
        "title": "Job Title",
        "company": "Company Name",
        "location": "City, Country",
        "date": "Month Year – Month Year",
        "description": "One sentence overall description (shown in italics)",
        "bullets": ["Achievement one", "Achievement two"]
      }}
    ],
    "education": [
      {{
        "degree": "Degree Title",
        "school": "School Name",
        "date": "Year – Year",
        "details": "Grade or extra info"
      }}
    ],
    "skills_and_interests": {{
      "languages": ["English (Native)", "French (Fluent)"],
      "tools": ["Excel (Advanced)", "Python (Intermediate)"],
      "competencies": ["Competency One", "Competency Two"],
      "interests": ["Interest One"]
    }}
  }},
  "keywords": [
    {{
      "word": "keyword phrase",
      "priority": 1,
      "score": 8,
      "matching_status": "missing"
    }}
  ],
  "proposals": [
    {{
      "id": "prop_1",
      "section": "experience",
      "entry_index": 0,
      "bullet_index": 1,
      "original": "Original sentence text",
      "proposed": "Rewritten sentence with keyword naturally included",
      "keywords": ["keyword phrase"]
    }}
  ],
  "redundancy_checks": [
    {{
      "id": "red_1",
      "location": "e.g. Profile summary and Amazon bullet 2",
      "issue": "The phrase 'cross-functional collaboration' is used 3 times across different sections.",
      "original": "The exact repeated sentence or phrase",
      "suggestion": "A cleaner, more varied alternative"
    }}
  ]
}}
"""
        try:
            response = client.models.generate_content(
                model='gemini-2.5-flash',
                contents=prompt,
                config=types.GenerateContentConfig(response_mime_type="application/json")
            )
            result = json.loads(response.text)
            return jsonify(result)
        except Exception as e:
            return jsonify({"error": f"Gemini API error: {e}"}), 502

    except Exception as e:
        return jsonify({"error": str(e)}), 500


# ── PDF builder ──────────────────────────────────────────────────────────────────

def _section_header_pdf(title, width, scale):
    style = ParagraphStyle(
        'SH', fontName='Helvetica-Bold',
        fontSize=round(10 * scale, 2), leading=round(12 * scale, 2),
        textColor=colors.HexColor('#1e293b'), spaceAfter=round(2 * scale, 2)
    )
    d = Drawing(width, round(3 * scale, 2))
    d.add(Line(0, round(1.5 * scale, 2), width, round(1.5 * scale, 2),
               strokeColor=colors.HexColor('#cbd5e1'), strokeWidth=0.75))
    return [Spacer(1, round(8 * scale, 2)), Paragraph(title, style), d, Spacer(1, round(4 * scale, 2))]


def _build_cv_buffer(cv_data, scale=1.0):
    """Render the CV to a BytesIO PDF at the given scale factor."""
    buf = BytesIO()
    margin   = max(0.36 * inch, 0.45 * inch * scale)
    usable_w = 8.5 * inch - 2 * margin

    doc = SimpleDocTemplate(buf, pagesize=letter,
                            leftMargin=margin, rightMargin=margin,
                            topMargin=margin, bottomMargin=margin)

    def fs(b): return round(b * scale, 2)
    def sp(b): return round(b * scale, 2)

    name_s    = ParagraphStyle('N',  fontName='Helvetica-Bold', fontSize=fs(19),   leading=sp(21),   textColor=colors.HexColor('#0f172a'), alignment=1)
    title_s   = ParagraphStyle('T',  fontName='Helvetica',      fontSize=fs(10.5), leading=sp(13),   textColor=colors.HexColor('#475569'), alignment=1, spaceAfter=sp(3))
    contact_s = ParagraphStyle('C',  fontName='Helvetica',      fontSize=fs(8.5),  leading=sp(10.5), textColor=colors.HexColor('#475569'), alignment=1, spaceAfter=sp(5))
    body_s    = ParagraphStyle('B',  fontName='Helvetica',      fontSize=fs(8.2),  leading=sp(11.5), textColor=colors.HexColor('#334155'))
    bullet_s  = ParagraphStyle('BU', parent=body_s, leftIndent=sp(11), firstLineIndent=sp(-9), spaceAfter=sp(2))
    jtitle_s  = ParagraphStyle('JT', fontName='Helvetica-Bold', fontSize=fs(8.8),  leading=sp(11),   textColor=colors.HexColor('#1e293b'))
    jmeta_s   = ParagraphStyle('JM', fontName='Helvetica',      fontSize=fs(8.2),  leading=sp(11),   textColor=colors.HexColor('#475569'))
    jdesc_s   = ParagraphStyle('JD', parent=body_s, fontName='Helvetica-Oblique',  textColor=colors.HexColor('#475569'), spaceAfter=sp(1.5))
    date_s    = ParagraphStyle('DA', fontName='Helvetica',      fontSize=fs(8.2),  leading=sp(11),   textColor=colors.HexColor('#475569'), alignment=2)

    story = []

    name    = strip_markdown(cv_data.get("name", ""))
    title   = strip_markdown(cv_data.get("title", ""))
    contact = strip_markdown(cv_data.get("contact", ""))
    profile = strip_markdown(cv_data.get("profile", ""))

    story.append(Paragraph(name, name_s))
    story.append(Paragraph(title.upper(), title_s))
    story.append(Paragraph(contact.replace(" | ", "   &bull;   "), contact_s))

    if profile:
        story.extend(_section_header_pdf("PROFILE", usable_w, scale))
        story.append(Paragraph(profile, body_s))

    experience = cv_data.get("experience", [])
    if experience:
        story.extend(_section_header_pdf("PROFESSIONAL EXPERIENCE", usable_w, scale))
        for idx, job in enumerate(experience):
            j_title   = strip_markdown(job.get("title", ""))
            j_company = strip_markdown(job.get("company", ""))
            j_loc     = strip_markdown(job.get("location", ""))
            j_date    = strip_markdown(job.get("date", ""))
            j_desc    = strip_markdown(job.get("description", ""))
            j_bullets = [strip_markdown(b) for b in job.get("bullets", [])]
            company_loc = j_company + (f" &bull; {j_loc}" if j_loc else "")

            row = Table([[Paragraph(j_title, jtitle_s), Paragraph(j_date, date_s)]],
                        colWidths=[usable_w * 0.7, usable_w * 0.3])
            row.setStyle(TableStyle([
                ('VALIGN',       (0,0),(-1,-1),'BOTTOM'),
                ('LEFTPADDING',  (0,0),(-1,-1), 0),
                ('RIGHTPADDING', (0,0),(-1,-1), 0),
                ('BOTTOMPADDING',(0,0),(-1,-1), 0),
                ('TOPPADDING',   (0,0),(-1,-1), 0),
            ]))
            story.append(row)
            story.append(Paragraph(company_loc, jmeta_s))
            story.append(Spacer(1, sp(1.5)))
            if j_desc:
                story.append(Paragraph(j_desc, jdesc_s))
                story.append(Spacer(1, sp(1.5)))
            for b in j_bullets:
                bt = b.strip()
                if not bt.startswith("&bull;") and not bt.startswith("•"):
                    bt = f"&bull; {bt}"
                story.append(Paragraph(bt, bullet_s))
            if idx < len(experience) - 1:
                story.append(Spacer(1, sp(4)))

    education = cv_data.get("education", [])
    if education:
        story.extend(_section_header_pdf("EDUCATION", usable_w, scale))
        for idx, ed in enumerate(education):
            e_deg  = strip_markdown(ed.get("degree", ""))
            e_sch  = strip_markdown(ed.get("school", ""))
            e_date = strip_markdown(ed.get("date", ""))
            e_det  = strip_markdown(ed.get("details", ""))
            row = Table([[Paragraph(e_deg, jtitle_s), Paragraph(e_date, date_s)]],
                        colWidths=[usable_w * 0.75, usable_w * 0.25])
            row.setStyle(TableStyle([
                ('VALIGN',       (0,0),(-1,-1),'BOTTOM'),
                ('LEFTPADDING',  (0,0),(-1,-1), 0),
                ('RIGHTPADDING', (0,0),(-1,-1), 0),
                ('BOTTOMPADDING',(0,0),(-1,-1), 0),
                ('TOPPADDING',   (0,0),(-1,-1), 0),
            ]))
            story.append(row)
            story.append(Paragraph(e_sch, jmeta_s))
            if e_det:
                story.append(Paragraph(e_det, body_s))
            if idx < len(education) - 1:
                story.append(Spacer(1, sp(3)))

    si = cv_data.get("skills_and_interests", {})
    if si:
        story.extend(_section_header_pdf("SKILLS & INTERESTS", usable_w, scale))
        lines = []
        if si.get("languages"):    lines.append(f"<b>Languages:</b> {', '.join(si['languages'])}")
        if si.get("tools"):        lines.append(f"<b>Tools:</b> {', '.join(si['tools'])}")
        if si.get("competencies"): lines.append(f"<b>Competencies:</b> {' &bull; '.join(si['competencies'])}")
        if si.get("interests"):    lines.append(f"<b>Interests:</b> {', '.join(si['interests'])}")
        story.append(Paragraph("<br/>".join(lines), body_s))

    doc.build(story)
    buf.seek(0)
    return buf


# ── Word builder ─────────────────────────────────────────────────────────────────

def _remove_table_borders(tbl):
    for cell in tbl.rows[0].cells:
        for border in ['top', 'bottom', 'left', 'right']:
            tc   = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBdr = OxmlElement('w:tcBdr')
            side  = OxmlElement(f'w:{border}')
            side.set(qn('w:val'), 'none')
            tcBdr.append(side)
            tcPr.append(tcBdr)


def _build_cv_docx(cv_data):
    """Build a Word (.docx) CV and return a seeked BytesIO buffer."""
    doc = Document()
    for section in doc.sections:
        section.top_margin    = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin   = Inches(0.6)
        section.right_margin  = Inches(0.6)

    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(9.5)

    def _para(text, bold=False, italic=False, size=9.5,
               align=WD_ALIGN_PARAGRAPH.LEFT, color=None,
               space_before=0, space_after=2):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after  = Pt(space_after)
        p.alignment = align
        run = p.add_run(text)
        run.bold = bold; run.italic = italic
        run.font.size = Pt(size)
        if color:
            run.font.color.rgb = RGBColor(*color)
        return p

    def _section_title(title):
        p = _para(title, bold=True, size=9.5, color=(30,41,59), space_before=6, space_after=1)
        # Add bottom border as section divider
        pPr  = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bot  = OxmlElement('w:bottom')
        bot.set(qn('w:val'), 'single'); bot.set(qn('w:sz'), '4')
        bot.set(qn('w:space'), '1');    bot.set(qn('w:color'), 'CBD5E1')
        pBdr.append(bot); pPr.append(pBdr)

    def _title_date_row(left_text, right_text, left_w=4.2, right_w=1.8):
        tbl = doc.add_table(rows=1, cols=2)
        tbl.autofit = False
        tbl.columns[0].width = Inches(left_w)
        tbl.columns[1].width = Inches(right_w)
        tbl.style = 'Table Grid'
        _remove_table_borders(tbl)
        r = tbl.rows[0]
        p0 = r.cells[0].paragraphs[0]
        p0.paragraph_format.space_after = Pt(0)
        run0 = p0.add_run(left_text)
        run0.bold = True; run0.font.size = Pt(9.5)
        run0.font.color.rgb = RGBColor(30,41,59)
        p1 = r.cells[1].paragraphs[0]
        p1.paragraph_format.space_after = Pt(0)
        p1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run1 = p1.add_run(right_text)
        run1.font.size = Pt(9)
        run1.font.color.rgb = RGBColor(71,85,105)

    # Header
    name    = strip_markdown(cv_data.get("name", ""))
    title   = strip_markdown(cv_data.get("title", ""))
    contact = strip_markdown(cv_data.get("contact", ""))
    profile = strip_markdown(cv_data.get("profile", ""))

    _para(name,          bold=True, size=18, align=WD_ALIGN_PARAGRAPH.CENTER, color=(15,23,42),  space_after=1)
    _para(title.upper(), size=10,            align=WD_ALIGN_PARAGRAPH.CENTER, color=(71,85,105), space_after=1)
    _para(contact.replace(" | ", "  •  "), size=8.5, align=WD_ALIGN_PARAGRAPH.CENTER, color=(71,85,105), space_after=4)

    if profile:
        _section_title("PROFILE")
        _para(profile, size=9, color=(51,65,85))

    experience = cv_data.get("experience", [])
    if experience:
        _section_title("PROFESSIONAL EXPERIENCE")
        for job in experience:
            j_title   = strip_markdown(job.get("title", ""))
            j_company = strip_markdown(job.get("company", ""))
            j_loc     = strip_markdown(job.get("location", ""))
            j_date    = strip_markdown(job.get("date", ""))
            j_desc    = strip_markdown(job.get("description", ""))
            j_bullets = [strip_markdown(b) for b in job.get("bullets", [])]

            _title_date_row(j_title, j_date)
            _para(j_company + (f" • {j_loc}" if j_loc else ""), size=9, color=(71,85,105), space_after=1)
            if j_desc:
                _para(j_desc, italic=True, size=9, color=(71,85,105), space_after=1)
            for b in j_bullets:
                p = doc.add_paragraph(style='List Bullet')
                p.paragraph_format.space_after = Pt(1)
                p.paragraph_format.left_indent = Inches(0.15)
                run = p.add_run(b.strip().lstrip('•').strip())
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(51,65,85)
            doc.add_paragraph().paragraph_format.space_after = Pt(2)

    education = cv_data.get("education", [])
    if education:
        _section_title("EDUCATION")
        for ed in education:
            e_deg  = strip_markdown(ed.get("degree", ""))
            e_sch  = strip_markdown(ed.get("school", ""))
            e_date = strip_markdown(ed.get("date", ""))
            e_det  = strip_markdown(ed.get("details", ""))
            _title_date_row(e_deg, e_date, left_w=4.5, right_w=1.5)
            _para(e_sch, size=9, color=(71,85,105), space_after=1)
            if e_det:
                _para(e_det, size=9, color=(51,65,85), space_after=2)

    si = cv_data.get("skills_and_interests", {})
    if si:
        _section_title("SKILLS & INTERESTS")
        rows = []
        if si.get("languages"):    rows.append(("Languages",    ', '.join(si['languages'])))
        if si.get("tools"):        rows.append(("Tools",        ', '.join(si['tools'])))
        if si.get("competencies"): rows.append(("Competencies", ' • '.join(si['competencies'])))
        if si.get("interests"):    rows.append(("Interests",    ', '.join(si['interests'])))
        for key, val in rows:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(1)
            rk = p.add_run(key + ': ')
            rk.bold = True; rk.font.size = Pt(9)
            rv = p.add_run(val)
            rv.font.size = Pt(9)
            rv.font.color.rgb = RGBColor(51,65,85)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ── Endpoints ────────────────────────────────────────────────────────────────────

@app.route('/api/generate-pdf', methods=['POST'])
def generate_pdf():
    try:
        cv_data = request.json
        if not cv_data:
            return jsonify({"error": "Missing CV JSON payload"}), 400
        scales    = [1.0, 0.97, 0.94, 0.91, 0.88, 0.85, 0.82, 0.79, 0.76]
        final_buf = None
        for scale in scales:
            candidate = _build_cv_buffer(cv_data, scale)
            if len(pypdf.PdfReader(candidate).pages) == 1:
                candidate.seek(0)
                final_buf = candidate
                break
        if final_buf is None:
            final_buf = _build_cv_buffer(cv_data, scales[-1])
            final_buf.seek(0)
        return send_file(final_buf, as_attachment=True,
                         download_name="Tailored_CV.pdf",
                         mimetype="application/pdf")
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route('/api/generate-docx', methods=['POST'])
def generate_docx():
    try:
        cv_data = request.json
        if not cv_data:
            return jsonify({"error": "Missing CV JSON payload"}), 400
        buf = _build_cv_docx(cv_data)
        return send_file(buf, as_attachment=True,
                         download_name="Tailored_CV.docx",
                         mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    except Exception as e:
        return jsonify({"error": str(e)}), 500


if __name__ == '__main__':
    app.run(debug=True, host='0.0.0.0', port=5000)
